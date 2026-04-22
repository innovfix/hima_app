from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('OTP Login Security Issue — Summary for Review', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

meta = doc.add_paragraph()
meta.add_run('Bug ID: ').bold = True
meta.add_run('BUG-001   ')
meta.add_run('Severity: ').bold = True
r = meta.add_run('S1 – Critical')
r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
meta.add_run('   Date: ').bold = True
meta.add_run('2026-04-21')

# 1. The Problem
doc.add_heading('1. The Problem (in simple terms)', level=1)
doc.add_paragraph(
    'Right now, the OTP entered during login is checked only inside the Android app, '
    'not on the server. Once the app thinks the OTP is correct, it calls the /login API '
    'with dummy values (code="0", code_verifier="0"), and the server accepts it without '
    'doing any real check.'
)
doc.add_paragraph(
    'A test OTP "011011" is also hard-coded inside the app itself. Anyone who opens the '
    'APK can see it.'
)

# 2. Why this is dangerous
doc.add_heading('2. Why this is dangerous', level=1)
points = [
    'Anyone with just a phone number can log in as that user — no SMS needed.',
    'A simple curl command returns a valid login token for any user.',
    'Full account takeover: chats, calls, wallet, profile — all accessible.',
    'The test OTP "011011" is visible in the release APK and works in production.',
    'Only Truecaller login is actually safe; everything else is bypassable.',
]
for p in points:
    doc.add_paragraph(p, style='List Bullet')

# 3. Proof
doc.add_heading('3. Proof', level=1)
doc.add_paragraph('A single request from any machine logs in as any user:')
code = doc.add_paragraph()
code_run = code.add_run(
    'curl -X POST https://demohima.himaapp.in/api/auth/login \\\n'
    '     -d "mobile=9876543210&code=0&code_verifier=0"\n\n'
    '→ HTTP 200 OK + valid JWT token'
)
code_run.font.name = 'Courier New'
code_run.font.size = Pt(10)

# 4. What needs to change
doc.add_heading('4. What needs to change', level=1)
fixes = [
    ('Server generates the OTP',
     'The server should create the OTP, send it via SMS, store it (with expiry), '
     'and verify it when the user logs in.'),
    ('Remove client-side OTP check',
     'The Android app should not compare OTPs locally. It should just send what the '
     'user typed to the server and let the server decide.'),
    ('Remove the hard-coded "011011" from the app',
     'No test OTP should ever live inside the APK. It must be removed from release builds.'),
    ('Keep a test OTP — but only on the server, only for test numbers',
     'For QA, the server can accept "011011" only for a small list of whitelisted test '
     'phone numbers, and only on dev/staging — never on production.'),
]
for title_text, desc in fixes:
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(title_text)
    r.bold = True
    doc.add_paragraph(desc)

# 5. What I need from you
doc.add_heading('5. What I need from you', level=1)
doc.add_paragraph(
    'To keep QA testing smooth while we fix this, I need approval and input on:'
)
asks = [
    'Approval to move OTP verification to the server side.',
    'Approval to remove the hard-coded "011011" from the Android app.',
    'A list of whitelisted test phone numbers that QA can use with the test OTP '
    '(e.g., 9999900001, 9999900002, 9999900003 …). Please confirm how many numbers '
    'you want to reserve for QA.',
    'Confirmation that the test OTP should work only on dev/staging, not on production.',
]
for a in asks:
    doc.add_paragraph(a, style='List Bullet')

# 6. Impact if we do nothing
doc.add_heading('6. If we do not fix this', level=1)
risks = [
    'Every user account is exposed — full takeover with just a phone number.',
    'Possible legal exposure (IT Act, DPDP Act) and Play Store policy violation.',
    'Financial loss if attackers drain wallets or misuse paid features.',
    'Serious damage to user trust if the issue is exploited or reported publicly.',
]
for r_ in risks:
    doc.add_paragraph(r_, style='List Bullet')

# Footer note
p = doc.add_paragraph()
p.add_run('\nPrepared for internal review. Please share the whitelisted test numbers '
          'and your decision so we can plan the fix.').italic = True

out = '/Users/apple/Desktop/hima_app_fresh/QA_Reports/BUG-001_OTP_Login_Security_Summary.docx'
doc.save(out)
print('Saved:', out)
