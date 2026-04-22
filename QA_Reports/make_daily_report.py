from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Base style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_status_run(paragraph, text, color_rgb):
    r = paragraph.add_run(text)
    r.bold = True
    r.font.color.rgb = color_rgb
    return r

GREEN = RGBColor(0x1F, 0x7A, 0x1F)
ORANGE = RGBColor(0xB0, 0x5A, 0x00)
RED = RGBColor(0xC0, 0x00, 0x00)
GREY = RGBColor(0x66, 0x66, 0x66)

# ---------- TITLE ----------
doc.add_heading('Daily QA & Security Work Report', level=0)

meta = doc.add_paragraph()
meta.add_run('Date: ').bold = True
meta.add_run('2026-04-21    ')
meta.add_run('Prepared by: ').bold = True
meta.add_run('Perumal    ')
meta.add_run('Branch: ').bold = True
meta.add_run('login_flow_test_21_04_2026    ')
meta.add_run('App: ').bold = True
meta.add_run('Hima (Android)')

# ---------- 1. SUMMARY ----------
doc.add_heading('1. Summary', level=1)
doc.add_paragraph(
    'Today I ran a full QA + security pass of the Hima Android login/signup flow '
    'across three testing tracks (API testing, automation, user-perspective '
    'testing), uncovered 16 bugs (4 S1 Critical, 4 S2 High, 5 S3 Medium, 3 S4 Low), '
    'and landed client-side fixes on branch login_flow_test_21_04_2026 for every '
    'bug that has a meaningful client-side mitigation. Bugs whose root cause is '
    'server-side are fully documented for the backend team.'
)
doc.add_paragraph(
    'Branch status: changes committed locally, not pushed to main — pending review.'
)

# ---------- 2. TESTING TRACKS ----------
doc.add_heading('2. Testing Performed', level=1)

doc.add_heading('2.1 API Testing', level=2)
doc.add_paragraph(
    'Black-box API testing against the dev backend using curl and Postman-style '
    'requests, independent of the Android client.'
)
for p in [
    'Exercised 8 auth endpoints: /first_install, /appsettings_list, /send_otp, '
    '/login, /register, /avatar_list, /userdetails, /delete_users.',
    'Confirmed /login accepts code="0" / code_verifier="0" and even random '
    'garbage, returning a valid JWT (BUG-001).',
    'Confirmed /userdetails returns any user profile when user_id is swapped in '
    'the body, regardless of JWT sub (BUG-002).',
    'Confirmed CORS reflects arbitrary origins with Allow-Credentials: true '
    '(BUG-003).',
    'Burst-tested /send_otp and /login — no rate limiting (BUG-004).',
    'Confirmed multiple 500-HTML error responses on invalid input (BUG-005), '
    'client-supplied OTP value (BUG-006), /register without OTP proof (BUG-007), '
    '90-day JWT lifetime (BUG-008), and additional medium/low defects.',
    'Total cases executed: 52. Artifact: Hima_Login_Flow_API_Test_Report_2026-04-21.pdf.',
]:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('2.2 Automation Testing', level=2)
doc.add_paragraph(
    'Automated regression + scripted scans to reproduce findings at scale and '
    'detect patterns across the codebase.'
)
for p in [
    'Ran the login-flow automated suite (see QA_Reports/automation/).',
    'Scripted user_id enumeration to confirm the IDOR (BUG-002) is exploitable '
    'at scale.',
    'Static scans across the Kotlin sources for hard-coded secrets, verbose '
    'logging interceptors, and raw PII Log.d calls (input for BUG-014, BUG-015).',
    'Artifact: Hima_Automated_Test_Report_2026-04-21.pdf.',
]:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('2.3 User-Perspective Testing', level=2)
doc.add_paragraph(
    'Hands-on testing on a real device, walking the app as an end user would, '
    'to catch issues a pure API test would miss.'
)
for p in [
    'Full manual login + OTP verification on a real mobile number.',
    'Verified resend-OTP cooldown (60s CountDownTimer) fires correctly on the '
    'device — relevant to BUG-004 client-side mitigation.',
    'Navigated home, wallet, and call flows to reproduce call-token / FCM-token '
    'and JWT logging (input for BUG-015).',
    'Monitored adb logcat throughout to identify sensitive data being written '
    'to device logs.',
    'Confirmed the existence of the "011011" backdoor OTP path via manual entry '
    '(input for BUG-001).',
]:
    doc.add_paragraph(p, style='List Bullet')

# ---------- 3. BUG STATUS TABLE ----------
doc.add_heading('3. Bug Status — All 16 Defects', level=1)
doc.add_paragraph(
    'Status legend: Fully fixed (client-side change complete, no backend dependency), '
    'Partially fixed (client-side mitigation landed, backend fix still required for '
    'full remediation), Documented (no meaningful client-side action possible — '
    'backend fix required).'
)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'ID'
hdr[1].text = 'Title'
hdr[2].text = 'Severity'
hdr[3].text = 'Status'

bug_rows = [
    ('BUG-001', 'OTP verified client-side only; /login accepts any code',
     'S1 Critical', 'Partially fixed'),
    ('BUG-002', 'IDOR on /userdetails — reads any user\'s profile',
     'S1 Critical', 'Documented'),
    ('BUG-003', 'CORS reflects arbitrary origins with credentials',
     'S1 Critical', 'Documented'),
    ('BUG-004', 'No rate limiting on /send_otp or /login',
     'S1 Critical', 'Partially fixed'),
    ('BUG-005', '/send_otp returns HTTP 500 HTML on invalid input',
     'S2 High', 'Documented'),
    ('BUG-006', 'OTP value is supplied by the client to /send_otp',
     'S2 High', 'Documented'),
    ('BUG-007', '/register requires no OTP / mobile-ownership proof',
     'S2 High', 'Documented'),
    ('BUG-008', 'JWT lifetime 90 days, no revocation, no /logout',
     'S2 High', 'Documented'),
    ('BUG-009', 'country_code accepts non-numeric strings on /send_otp',
     'S3 Medium', 'Documented'),
    ('BUG-010', 'HTTP 200 returned on business-logic failures',
     'S3 Medium', 'Documented'),
    ('BUG-011', '/first_install leaks caller IP in response',
     'S3 Medium', 'Documented'),
    ('BUG-012', 'Referral code not validated on /register',
     'S3 Medium', 'Documented'),
    ('BUG-013', 'No size cap on describe_yourself field',
     'S3 Medium', 'Documented'),
    ('BUG-014', 'Cashfree test credentials hard-coded in Kotlin source',
     'S4 Low', 'Partially fixed'),
    ('BUG-015', 'Heavy Log.d of OTP / request / response in release builds',
     'S4 Low', 'Fully fixed'),
    ('BUG-016', 'delete_users error message is cryptic catch-all',
     'S4 Low', 'Documented'),
]
for row in bug_rows:
    cells = table.add_row().cells
    cells[0].text = row[0]
    cells[1].text = row[1]
    cells[2].text = row[2]
    cells[3].text = row[3]

# Count summary line
summary_p = doc.add_paragraph()
summary_p.add_run('\nTotals: ').bold = True
summary_p.add_run('1 fully fixed, 3 partially fixed, 12 documented for backend fix.')

# ---------- 4. CLIENT-SIDE CHANGES LANDED ----------
doc.add_heading('4. Client-Side Changes Landed Today', level=1)

# 4.1 BUG-015
doc.add_heading('4.1 BUG-015 — Release-build logging (Fully fixed)', level=2)
for p in [
    'Gated main Retrofit HttpLoggingInterceptor on BuildConfig.DEBUG in '
    'AppModule.kt — release builds log at Level.NONE.',
    'Gated the custom FCM request/response debug interceptor on BuildConfig.DEBUG '
    '(it previously printed the full JWT).',
    'Removed or redacted 14 sensitive Log.d calls across 11 files — OTPs, JWTs, '
    'FCM tokens, Agora tokens, raw SMS, and raw mobile numbers no longer appear '
    'in logs.',
    'Added R8 -assumenosideeffects rule in proguard-rules.pro so Log.v/d/i are '
    'stripped from release once minification is enabled; Log.w / Log.e preserved.',
]:
    doc.add_paragraph(p, style='List Bullet')

# 4.2 BUG-001 partial
doc.add_heading('4.2 BUG-001 — OTP verification (Partially fixed)', level=2)
for p in [
    'Removed the hard-coded "011011" backdoor OTP from VerifyOTPActivity.kt:143 '
    'and NewLoginActivity.kt:835 — the APK no longer contains the string.',
    'Removed the disjunct (enteredOTP == "011011") from both OTP-compare blocks.',
]:
    doc.add_paragraph(p, style='List Bullet')
pp = doc.add_paragraph()
pp.add_run('Backend still required: ').bold = True
pp.add_run(
    'server must generate the OTP, store it with expiry, and validate it on '
    '/login instead of accepting code="0". Client still performs a local compare '
    'only as a UX check — this is not authentication.'
)

# 4.3 BUG-014 partial
doc.add_heading('4.3 BUG-014 — Cashfree keys in source (Partially fixed)', level=2)
for p in [
    'Removed the hard-coded x-client-id / x-client-secret headers from '
    'CashfreeApiService.kt — the keys no longer exist in any Kotlin file.',
    'Added CASHFREE_CLIENT_ID and CASHFREE_CLIENT_SECRET to local.properties '
    '(already gitignored).',
    'Wired local.properties values into BuildConfig via app/build.gradle.kts.',
    'Updated CashfreeRetrofitClient.kt to inject the credentials through an '
    'OkHttp interceptor that reads BuildConfig at runtime.',
]:
    doc.add_paragraph(p, style='List Bullet')
pp = doc.add_paragraph()
pp.add_run('Still required: ').bold = True
pp.add_run(
    'the keys that were already committed to git history must be rotated in the '
    'Cashfree dashboard — the leaked values should be treated as burned. '
    'Before production launch, PAN verification should move to the backend; '
    'BuildConfig still bakes the secret into the APK.'
)

# 4.4 BUG-004 partial
doc.add_heading('4.4 BUG-004 — Rate limiting (Partially fixed)', level=2)
for p in [
    'Verified the 60-second CountDownTimer gating the Resend-OTP button is '
    'enforced in both VerifyOTPActivity.kt (startTimer at line 160) and '
    'NewLoginActivity.kt (startTimer at line 850). This prevents in-app rapid '
    'resend abuse.',
]:
    doc.add_paragraph(p, style='List Bullet')
pp = doc.add_paragraph()
pp.add_run('Backend still required: ').bold = True
pp.add_run(
    'attackers can skip the app and call /send_otp directly via curl — only '
    'server-side per-mobile / per-IP rate limiting truly closes this. Also '
    'needed: per-mobile attempt cap on /login once OTP moves server-side.'
)

# ---------- 5. DOCUMENTED-ONLY BUGS ----------
doc.add_heading('5. Documented-Only Bugs (backend fix required)', level=1)
doc.add_paragraph(
    'The following bugs have no meaningful client-side mitigation — the root '
    'cause is on the server. They are captured in the API test report with '
    'repro steps, proofs, and recommended fixes for the backend team.'
)

documented = [
    ('BUG-002', 'Server must derive user id from JWT sub; stop trusting body user_id.'),
    ('BUG-003', 'Hard-code allowed origins; disable Access-Control-Allow-Credentials '
                'for the API origin.'),
    ('BUG-005', 'Add FormRequest validators; return JSON 400, not 500 HTML.'),
    ('BUG-006', 'Server generates the OTP cryptographically; remove otp param '
                'from /send_otp payload.'),
    ('BUG-007', '/register must require a server-issued OTP token that proves '
                'mobile ownership.'),
    ('BUG-008', 'Short-lived access token + refresh rotation; add /logout that '
                'blacklists jti.'),
    ('BUG-009', 'Strict integer validation on country_code.'),
    ('BUG-010', 'Return 4xx for validation / business failures, not 200 + success:false.'),
    ('BUG-011', 'Remove caller IP from /first_install response body.'),
    ('BUG-012', 'Validate referral code on /register — reject unknown codes or '
                'explicitly attribute.'),
    ('BUG-013', 'Enforce per-field max length at the framework layer.'),
    ('BUG-016', 'Replace generic "please mail your mobile number" message with '
                'a specific one.'),
]
for bug_id, fix in documented:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f'{bug_id}: ')
    r.bold = True
    p.add_run(fix)

# ---------- 6. FILES CHANGED ----------
doc.add_heading('6. Files Changed on Branch', level=1)
files_changed = [
    ('app/build.gradle.kts', 'BuildConfig fields for Cashfree keys.'),
    ('local.properties', 'Cashfree keys (gitignored).'),
    ('app/proguard-rules.pro', 'R8 -assumenosideeffects for Log.v/d/i.'),
    ('app/src/main/java/com/gmwapp/hima/dagger/AppModule.kt',
     'HttpLoggingInterceptor + custom FCM debug gated on BuildConfig.DEBUG.'),
    ('app/src/main/java/com/gmwapp/hima/verification/CashfreeApiService.kt',
     'Removed hard-coded x-client-id / x-client-secret headers.'),
    ('app/src/main/java/com/gmwapp/hima/verification/CashfreeRetrofitClient.kt',
     'Inject Cashfree keys via OkHttp interceptor reading BuildConfig.'),
    ('app/src/main/java/com/gmwapp/hima/activities/VerifyOTPActivity.kt',
     'Removed "011011" backdoor; scrubbed OTP Log.d.'),
    ('app/src/main/java/com/gmwapp/hima/activities/NewLoginActivity.kt',
     'Removed "011011" backdoor; scrubbed OTP / token / SMS / mobile Log.d.'),
    ('app/src/main/java/com/gmwapp/hima/activities/WalletActivity.kt',
     'Removed JWT Log.d.'),
    ('app/src/main/java/com/gmwapp/hima/activities/MainActivity.kt',
     'Removed FCM device-token Log.d.'),
    ('app/src/main/java/com/gmwapp/hima/activities/SelectLanguageActivity.kt',
     'Removed mobile-number Log.d.'),
    ('app/src/main/java/com/gmwapp/hima/agora/MyFirebaseMessagingService.kt',
     'Redacted new-FCM-token Log.d.'),
    ('app/src/main/java/com/gmwapp/hima/agora/female/FemaleAudioCallingActivity.kt',
     'Redacted Agora-token Log.d.'),
    ('app/src/main/java/com/gmwapp/hima/agora/female/FemaleVideoCallingActivity.kt',
     'Redacted Agora-token Log.d.'),
    ('app/src/main/java/com/gmwapp/hima/agora/male/MaleAudioCallingActivity.kt',
     'Redacted Agora-token Log.d.'),
    ('app/src/main/java/com/gmwapp/hima/viewmodels/AgoraViewModel.kt',
     'Redacted Agora-token Log.d.'),
    ('app/src/main/java/com/gmwapp/hima/viewmodels/FcmTokenViewModel.kt',
     'Removed FCM-response Log.d.'),
    ('app/src/main/java/com/gmwapp/hima/viewmodels/LoginViewModel.kt',
     'Removed full-response-body Log.d (login + sendOTP).'),
]
ftable = doc.add_table(rows=1, cols=2)
ftable.style = 'Light Grid Accent 1'
ftable.rows[0].cells[0].text = 'File'
ftable.rows[0].cells[1].text = 'Change'
for path, change in files_changed:
    cells = ftable.add_row().cells
    cells[0].text = path
    cells[1].text = change

# ---------- 7. WORK ATTRIBUTION ----------
doc.add_heading('7. Work Attribution — What I Did vs What Claude Assisted With', level=1)
doc.add_paragraph(
    'Being transparent about tooling: AI assistance (Claude Code) was used for '
    'code analysis, bug explanation in simple language, and executing edits '
    'under my direction. All testing, judgment calls, and the QA report itself '
    'are my own.'
)

doc.add_heading('7.1 Done by me (Perumal)', level=2)
for p in [
    'Ran all three testing tracks end-to-end: API testing with curl / Postman, '
    'automation runs, and user-perspective testing on a real device.',
    'Captured proof responses, logcat dumps, and curl transcripts as evidence.',
    'Identified and triaged all 16 bugs; assigned severities; wrote the full '
    'Login/Signup Flow API Test Report (52 cases).',
    'Decided what to fix today, what to escalate to backend, and what to defer.',
    'Reviewed all code changes before they were committed to the branch.',
    'Chose to keep the branch local (no push to main) pending senior review.',
]:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('7.2 Used Claude for', level=2)
for p in [
    'Plain-language explanations of each bug (impact, root cause, severity '
    'rationale) to share with non-technical stakeholders.',
    'Code analysis to locate every Log.d call printing sensitive data and every '
    'hard-coded secret in the codebase.',
    'Implementing the code edits for BUG-001, BUG-014, and BUG-015 under my '
    'direction (remove backdoor, move keys to BuildConfig, gate logging).',
    'Drafting the senior-review docx for BUG-001 and this daily report.',
    'Cross-checking that nothing was missed (e.g. the second copy of the '
    '"011011" backdoor in NewLoginActivity.kt).',
]:
    doc.add_paragraph(p, style='List Bullet')

# ---------- 8. ASKS / NEXT STEPS ----------
doc.add_heading('8. Pending / Asks for Senior & Backend', level=1)
for p in [
    'Senior approval to move OTP verification to the server side (BUG-001 full fix).',
    'Whitelisted test phone numbers for QA to use once server-side OTP is in place.',
    'Backend team pickup of the 12 documented-only bugs (BUG-002, 003, 005, 006, '
    '007, 008, 009, 010, 011, 012, 013, 016).',
    'Rotate the Cashfree test credentials in the Cashfree dashboard (BUG-014) '
    'since the leaked values remain in git history and must be considered burned.',
    'Separate QA-tested task to enable isMinifyEnabled = true so the ProGuard '
    'log stripping activates in release (BUG-015 defence-in-depth).',
    'Approval to push login_flow_test_21_04_2026 to the remote and open a PR.',
]:
    doc.add_paragraph(p, style='List Bullet')

# ---------- 9. ARTIFACTS ----------
doc.add_heading('9. Artifacts', level=1)
for p in [
    'QA_Reports/login_flow_api_test_report_2026-04-21.md — full API test report '
    '(52 cases, 16 bugs).',
    'QA_Reports/Hima_Login_Flow_API_Test_Report_2026-04-21.pdf — PDF export.',
    'QA_Reports/Hima_Automated_Test_Report_2026-04-21.pdf — automation run.',
    'QA_Reports/automation/ — automation source + latest results JSON.',
    'QA_Reports/BUG-001_OTP_Login_Security_Summary.docx — senior-review brief.',
    'QA_Reports/Daily_Report_2026-04-21.docx — this report.',
]:
    doc.add_paragraph(p, style='List Bullet')

p = doc.add_paragraph()
p.add_run('\nEnd of report — 2026-04-21').italic = True

out = '/Users/apple/Desktop/hima_app_fresh/QA_Reports/Daily_Report_2026-04-21.docx'
doc.save(out)
print('Saved:', out)
